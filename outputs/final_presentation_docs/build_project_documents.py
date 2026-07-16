from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path(__file__).resolve().parent
PROJECT_TITLE = "Academic Timetable Scheduling System"
TEAM = "Team: Roy, Xian Yang, Kai Xian, Christie, Ikin, Anastasia"
TEAM_NUMBER = "Team number: To be confirmed"
SUPERVISOR = "Supervisor: Ms Yang"
COURSE = "INF1009 ITP Final Presentation - 27 July 2026"


def set_font(run, name="Calibri", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text) if text is not None else "")
    set_font(run, size=9.5, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def table_geometry(table, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)
    for child in list(tbl_grid):
        tbl_grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        tbl_grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")

            tc_mar = tc_pr.find(qn("w:tcMar"))
            if tc_mar is None:
                tc_mar = OxmlElement("w:tcMar")
                tc_pr.append(tc_mar)
            for side, value in [("top", 80), ("bottom", 80), ("start", 120), ("end", 120)]:
                node = tc_mar.find(qn(f"w:{side}"))
                if node is None:
                    node = OxmlElement(f"w:{side}")
                    tc_mar.append(node)
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")


def make_doc(title, subtitle):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(2)
    r = title_p.add_run(title)
    set_font(r, size=21, bold=True, color="0B2545")

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(10)
    r = sub.add_run(subtitle)
    set_font(r, size=10.5, color="555555")

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(12)
    for idx, item in enumerate([COURSE, TEAM_NUMBER, SUPERVISOR, TEAM]):
        run = meta.add_run(item)
        set_font(run, size=9.5, bold=idx == 0, color="1F3A5F" if idx == 0 else "555555")
        if idx < 3:
            meta.add_run("\n")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(PROJECT_TITLE)
    set_font(run, size=8.5, color="777777")
    return doc


def add_para(doc, text, bold_start=None):
    p = doc.add_paragraph()
    if bold_start and text.startswith(bold_start):
        r = p.add_run(bold_start)
        set_font(r, bold=True)
        r = p.add_run(text[len(bold_start) :])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_font(r, size=10.2)


def add_numbers(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_font(r, size=10.2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table_geometry(table, widths)
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], "E8EEF5")
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color="0B2545")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    table_geometry(table, widths)
    doc.add_paragraph()
    return table


def add_callout(doc, label, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table_geometry(table, [9360])
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(label + ": ")
    set_font(r, size=10.2, bold=True, color="1F3A5F")
    r = p.add_run(body)
    set_font(r, size=10.2)
    doc.add_paragraph()


def technical_breakdown():
    doc = make_doc(
        "Full Technical Breakdown",
        "Architecture, implementation, constraints, validation, export, testing, and future work.",
    )

    doc.add_heading("1. Executive Summary", level=1)
    add_para(
        doc,
        "The Academic Timetable Scheduling System is a full-stack prototype for importing academic timetable requirements, "
        "validating the data, generating an optimized timetable, reviewing hard and soft conflicts, applying repairs, and exporting "
        "the final output in a system-ready CSV/XLSX format.",
    )
    add_callout(
        doc,
        "Technical thesis",
        "The project is strongest when presented as an end-to-end scheduling pipeline: Excel ingestion -> relational data model -> validation -> CP-SAT optimization -> conflict review -> export.",
    )

    doc.add_heading("2. Problem and Objectives", level=1)
    add_bullets(
        doc,
        [
            "Manual timetable planning must coordinate modules, student groups, lecturers, rooms, class types, delivery modes, week patterns, and fixed sessions.",
            "Errors such as room double-booking, staff clashes, invalid online/physical placements, or capacity mismatches are difficult to detect manually.",
            "The project objective is to reduce manual work by creating an import-to-export workflow that validates data early and makes conflicts visible.",
            "The success criteria are: import realistic Excel data, generate a feasible timetable when possible, show reviewable conflicts when input is over-constrained, and export a downstream-compatible template.",
        ],
    )

    doc.add_heading("3. Technology Stack", level=1)
    add_table(
        doc,
        ["Layer", "Technology", "Purpose"],
        [
            ["Frontend", "React, TypeScript, Vite", "Single-page workflow for import, generation, review, conflict repair, and export."],
            ["Backend API", "FastAPI", "REST endpoints for upload, validation, scheduling, database, review, and export."],
            ["Persistence", "SQLAlchemy, SQLite", "Local relational model for programmes, modules, staff, rooms, sessions, runs, and violations."],
            ["Excel Processing", "pandas, openpyxl", "Read uploaded workbooks and write final CSV/XLSX outputs."],
            ["Optimization", "Google OR-Tools CP-SAT", "Constraint programming model for assigning sessions to rooms and time slots."],
            ["Testing", "pytest, Ruff, TypeScript build", "Regression tests, style checks, and production frontend build verification."],
        ],
        [1450, 2150, 5760],
    )

    doc.add_heading("4. End-to-End Workflow", level=1)
    add_numbers(
        doc,
        [
            "Admin uploads one or more Excel requirement templates.",
            "Import service normalizes columns, merges optional remarks, applies defaults, and builds preview rows.",
            "Requirement validation checks mandatory fields, references, delivery mode, room feasibility, staff IDs, week patterns, fixed timings, and duplicate requirement IDs.",
            "Clean data is saved into SQLite tables and previous schedule runs are cleared.",
            "Soft priority weights are read from the settings table and passed to the solver.",
            "CP-SAT generates assignments for each session to exactly one compatible room and time slot.",
            "Constraint service checks the generated schedule and stores hard/soft violations for review.",
            "React review screens display timetable views, filters, conflict tables, quick fixes, and manual move tools.",
            "Export service produces system-template CSV/XLSX files for submission or downstream upload.",
        ],
    )

    doc.add_heading("5. Backend Architecture", level=1)
    add_table(
        doc,
        ["Component", "Key Files", "Responsibility"],
        [
            ["App entrypoint", "backend/app/main.py", "Registers FastAPI app, CORS, health check, database startup, and routers."],
            ["Routes", "backend/app/routes/*.py", "Expose upload, validation, data, database, schedule, soft priority, and export APIs."],
            ["Models", "backend/app/models/*.py", "Define programmes, modules, staff, rooms, groups, sessions, runs, scheduled sessions, and violations."],
            ["Import", "services/import_service.py, requirement_input_service.py", "Read Excel, normalize columns, validate rows, create/update requirement sessions."],
            ["Scheduling", "services/schedule_service.py, solver/*.py", "Build model, run CP-SAT, persist assignments, and run post-generation checks."],
            ["Conflict repair", "services/constraint_service.py, quick_fix_service.py", "Detect stored issues and propose clean room/time alternatives."],
            ["Export", "services/export_service.py", "Map internal scheduled rows back to the required system upload template shape."],
        ],
        [1450, 3150, 4760],
    )

    doc.add_heading("6. Data Model", level=1)
    add_para(doc, "The core database tables separate stable reference data from imported requirements and generated timetable results.")
    add_table(
        doc,
        ["Table", "Main Data Stored"],
        [
            ["programmes", "Programme code, name, and academic year context."],
            ["modules", "Module code, title, host key, and term."],
            ["student_groups", "Group code, programme, year, and size."],
            ["staff", "Staff ID, staff name, and host key."],
            ["rooms", "Room code, type, capacity, campus mode, virtual flag, and recording support."],
            ["time_slots", "Day, start/end time, duration, and week pattern."],
            ["sessions", "Imported timetable requirements, staff links, class size, delivery mode, venue needs, week pattern, and preferences."],
            ["schedule_runs", "Solver status, run status, soft score, hard violation count, and message."],
            ["scheduled_sessions", "Final room/time assignment for each scheduled requirement."],
            ["constraint_violations", "Stored hard/soft issues with affected session IDs and explanation messages."],
        ],
        [2200, 7160],
    )

    doc.add_heading("7. Import and Validation Pipeline", level=1)
    add_para(
        doc,
        "The importer reads an Input_Template sheet when present and joins optional data from Remarks_(optional). It also tolerates column aliases so minor naming differences in Excel files do not break the workflow.",
    )
    add_bullets(
        doc,
        [
            "Canonical columns include Requirement ID, Programme, Year, Module Code, Class Type, Duration, Delivery Mode, Venue Type Required, Exact Class Size, Staff IDs, week fields, fixed time fields, preferences, and remarks.",
            "Bad uploads are rejected as a batch. This prevents half-imported data from corrupting the scheduling state.",
            "Preview rows are returned to the frontend so users can correct validation errors in an editable grid.",
            "Edited rows are revalidated through the same backend rules before replacing the real requirements table.",
        ],
    )

    doc.add_heading("8. CP-SAT Solver Design", level=1)
    add_para(
        doc,
        "The solver creates a boolean variable for every compatible assignment candidate. A variable x(session, time_slot, room) is 1 when that session is placed in that exact room and time.",
    )
    add_table(
        doc,
        ["Solver Element", "Implementation Meaning"],
        [
            ["Decision variable", "x[session_id, timeslot_id, room_id] indicates one possible placement."],
            ["Exactly-once rule", "Every requirement must choose exactly one compatible placement."],
            ["Candidate filtering", "Invalid rooms/times are removed before model creation using scheduling_rules.py."],
            ["Objective", "Minimize weighted soft penalties while satisfying hard constraints."],
            ["Strict-first mode", "Try hard room/staff/group no-overlap constraints first."],
            ["Relaxed fallback", "If strict solving is infeasible or times out, rebuild with heavy clash penalties so conflicts can be reviewed."],
        ],
        [2400, 6960],
    )

    doc.add_heading("9. Constraint Model", level=1)
    add_table(
        doc,
        ["Constraint Type", "Examples", "System Behavior"],
        [
            ["Hard", "Room double-booking, staff double-booking, student group overlap.", "Blocked in strict solve or heavily penalized in reviewable fallback."],
            ["Hard", "Room capacity, delivery-room mismatch, venue type mismatch.", "Invalid candidate assignments are filtered or reported as hard issues."],
            ["Hard", "Fixed sessions must use fixed day/start/end.", "Validation and post-solve checks detect violations."],
            ["Soft", "Preferred day mismatch, soft avoid day placement, online not Mon/Tue.", "Converted into weighted objective penalties."],
            ["Soft", "Tutor idle gap, short campus day, long consecutive student day, online/F2F adjacent switch.", "Scored and stored as review warnings after generation."],
        ],
        [1500, 3300, 4560],
    )

    doc.add_heading("10. Solver Performance Improvement", level=1)
    add_para(
        doc,
        "A key improvement was replacing expensive pairwise resource clash checks with bucketed no-overlap constraints. This reduced model-building pressure on large imports while preserving reviewable hard-conflict behavior.",
    )
    add_bullets(
        doc,
        [
            "Large input stress test: 116 sessions across 20 programmes, 255 time slots, and 109 rooms.",
            "Naive candidate upper bound was about 3.2 million combinations; filtered candidates were about 98,940.",
            "Earlier pairwise clash loops could reach hundreds of millions of comparisons across room, staff, and group buckets.",
            "The optimized approach groups variables by resource and conflicting time bucket, adding aggregate constraints instead of comparing every pair.",
        ],
    )

    doc.add_heading("11. Frontend Workflow", level=1)
    add_table(
        doc,
        ["Page", "Purpose"],
        [
            ["Dashboard", "Shows current data state, validation state, and latest schedule state."],
            ["Import Data", "Uploads Excel files, shows import summary, editable preview grid, and generation readiness."],
            ["Database", "Allows admin management of rooms, staff, programmes, modules, and student groups."],
            ["Generate Timetable", "Runs CP-SAT generation using configured soft priorities."],
            ["Review Timetable", "Shows timetable, filters, versions, conflicts, explanations, manual moves, and quick fixes."],
            ["Export", "Downloads CSV/XLSX once hard conflicts are resolved."],
            ["Settings", "Manages soft constraint priority ranking."],
        ],
        [2200, 7160],
    )

    doc.add_heading("12. Conflict Review and Quick Fixes", level=1)
    add_para(
        doc,
        "The review page turns solver results into an operational interface. It separates raw issues from modules that need reassignment, highlights hard conflicts, and provides suggested clash-free fixes.",
    )
    add_bullets(
        doc,
        [
            "Manual move controls let users adjust day, time, room, and staff placement.",
            "The timetable grid highlights available, soft-warning, and blocked slots for the selected session.",
            "QuickFixService ranks candidate fixes as venue change, time change, or alternative best placement.",
            "After a move or quick fix, the schedule is refreshed and constraints are rechecked.",
        ],
    )

    doc.add_heading("13. Export Format", level=1)
    add_para(
        doc,
        "The export maps internal scheduled rows into the required system-template columns, so the final output resembles the expected downstream upload format rather than raw backend IDs.",
    )
    add_table(
        doc,
        ["Export Column", "Mapping Logic"],
        [
            ["Module", "Module code from the scheduled session."],
            ["Class Type", "Lecture, Tutorial, Laboratory, Online, Workshop, etc."],
            ["Template", "Grouped template number based on module, class type, staff, time, and teaching weeks."],
            ["Group", "All, L1, T1, W1, or generated group label depending on class type and cohort coverage."],
            ["Day / Start / End", "Short day label and compact 24-hour time strings such as Mon, 0900, 1100."],
            ["Room1 / Staff1 / Staff2", "Assigned room and uppercase staff names."],
            ["Tri Week", "Custom weeks or start-end week pattern converted into upload-compatible values."],
            ["Recording Mode", "A0 when the selected room supports recording."],
            ["Remark", "Remarks or combined-module notation such as w ESE1109."],
        ],
        [2300, 7060],
    )

    doc.add_heading("14. Testing and Verification", level=1)
    add_bullets(
        doc,
        [
            "Backend test coverage includes upload routes, import validation, solver behavior, soft priorities, constraints, manual moves, database routes, session CRUD, and exports.",
            "Regression tests cover important edge cases such as infeasible capacity, fixed-session clashes, duplicate requirement IDs, edited import rows, and reviewable hard conflicts.",
            "Frontend verification includes TypeScript production build and targeted linting of touched files.",
            "Known full frontend lint issues remain in older files, but touched files passed targeted lint checks according to project notes.",
        ],
    )

    doc.add_heading("15. Current Limitations", level=1)
    add_bullets(
        doc,
        [
            "Authentication and user roles are not implemented.",
            "Uploaded imports replace current sessions and generated runs.",
            "Advanced shared/common-module merging is limited and reserved for future work.",
            "Custom week overlap logic can be improved for more complex academic calendars.",
            "Solver results may vary between runs because multi-worker CP-SAT search is not fully deterministic.",
        ],
    )

    doc.add_heading("16. Future Work", level=1)
    add_bullets(
        doc,
        [
            "Add authentication, role-based permissions, and audit history.",
            "Improve support for shared/common modules and combined cohort sessions.",
            "Add deterministic solver options such as fixed random seed, single worker mode, and stronger tie-break penalties.",
            "Integrate staff/room unavailability constraints and richer calendar rules.",
            "Improve export compatibility with additional downstream templates.",
        ],
    )
    return doc


def presentation_and_poster():
    doc = make_doc(
        "Presentation and Poster Content",
        "Slide-by-slide content, demo flow, Q&A preparation, and poster-ready wording.",
    )

    doc.add_heading("1. Presentation Goal", level=1)
    add_para(
        doc,
        "The presentation should show that the team built a working technical system, not only a UI mock-up. The strongest story is the complete pipeline from messy Excel requirements to validated, optimized, reviewable, and export-ready timetables.",
    )

    doc.add_heading("2. 20-Minute Slide Plan", level=1)
    add_table(
        doc,
        ["Slide", "Owner", "Time", "Main Message", "Content"],
        [
            ["1", "Roy", "0:30", "Introduce project.", "Title, team members, supervisor Ms Yang, project name, and one-line purpose."],
            ["2", "Roy", "1:20", "Manual scheduling is constraint-heavy.", "Explain staff, rooms, class size, delivery mode, fixed sessions, and student group clashes."],
            ["3", "Roy", "1:20", "Define objectives.", "Import Excel data, validate errors, generate timetable, review conflicts, export system-ready output."],
            ["4", "Roy", "0:50", "Show architecture overview.", "Frontend -> FastAPI -> SQLite -> CP-SAT -> review/export workflow."],
            ["5", "Xian Yang", "1:40", "Data is normalized before scheduling.", "Input_Template, optional remarks, canonical columns, batch validation."],
            ["6", "Xian Yang", "1:40", "Validation protects the solver.", "Required fields, staff IDs, rooms, capacity, delivery mode, fixed times, week patterns."],
            ["7", "Kai Xian", "1:50", "Solver uses assignment variables.", "x(session, room, time slot), candidate filtering, exactly-once assignment."],
            ["8", "Kai Xian", "1:35", "Hard and soft constraints are separated.", "Hard: room/staff/group/capacity/fixed. Soft: preferences, gaps, online/F2F switches."],
            ["9", "Kai Xian", "1:15", "Optimization was improved.", "Strict-first solving, relaxed fallback, bucketed constraints, robust 116-row stress test."],
            ["10", "Christie", "1:30", "Frontend supports the admin workflow.", "Dashboard, Import Data, Database, Generate, Review, Export, Settings."],
            ["11", "Christie", "1:35", "Import preview improves recoverability.", "Editable grid, highlighted cells, issues-only filter, validate-and-apply edits."],
            ["12", "Ikin", "1:40", "Review screen turns schedules into decisions.", "Timetable grid, filters, schedule versions, explanations, conflict list."],
            ["13", "Ikin", "1:50", "Conflicts can be repaired.", "Manual moves, blocked/available slots, quick fix suggestions, recheck after move."],
            ["14", "Anastasia", "1:45", "Export and testing complete the workflow.", "System-template CSV/XLSX output, backend tests, build checks."],
            ["15", "Anastasia", "1:30", "Impact and future work.", "Less manual work, fewer errors, future auth, better shared modules, deterministic solving."],
        ],
        [650, 1250, 800, 2200, 4460],
    )

    doc.add_heading("3. Suggested Slide Content", level=1)
    slide_content = [
        ("Slide 1 - Title", ["Academic Timetable Scheduling System", "Team: Roy, Xian Yang, Kai Xian, Christie, Ikin, Anastasia", "Supervisor: Ms Yang"]),
        ("Slide 2 - Background", ["Timetable planning requires matching modules, lecturers, student groups, rooms, and time slots.", "Manual planning can miss clashes and takes repeated checking.", "A practical scheduler must support both hard rules and softer preferences."]),
        ("Slide 3 - Problem Statement", ["How can we transform Excel timetable requirements into a validated and export-ready timetable while reducing manual clash checking?", "The system must detect invalid input early, generate feasible schedules, and make remaining issues easy to resolve."]),
        ("Slide 4 - System Architecture", ["React frontend for workflow screens.", "FastAPI backend for API and business logic.", "SQLite database for local persistence.", "OR-Tools CP-SAT solver for scheduling optimization.", "pandas/openpyxl for Excel import/export."]),
        ("Slide 5 - Data Import", ["Reads Input_Template and optional Remarks_(optional).", "Normalizes column aliases into canonical fields.", "Creates preview rows for frontend correction.", "Batch import avoids partial corrupted data."]),
        ("Slide 6 - Validation", ["Checks required fields and duplicate requirement IDs.", "Verifies programme, staff, group, room, class size, delivery mode, week pattern, and fixed time compatibility.", "Invalid rows are highlighted for correction before saving."]),
        ("Slide 7 - Solver Model", ["Each boolean variable represents assigning one session to one room/time slot.", "Every session must be scheduled exactly once.", "Candidate filtering reduces impossible assignments before optimization."]),
        ("Slide 8 - Constraints", ["Hard constraints prevent resource clashes and invalid placements.", "Soft constraints rank timetable quality without blocking generation.", "Weighted priorities let users influence timetable preferences."]),
        ("Slide 9 - Performance and Robustness", ["Bucketed no-overlap constraints reduce expensive pairwise comparisons.", "Strict solve is attempted first.", "Relaxed fallback keeps over-constrained inputs reviewable.", "Stress sample covers 20 programmes and 116 sessions."]),
        ("Slide 10 - Frontend Workflow", ["Pages follow the real admin workflow: import, database, generate, review, export.", "Session state preserves context while navigating.", "Inline status cards and progress indicators guide each step."]),
        ("Slide 11 - Editable Import Preview", ["Failed uploads show row-level preview.", "Issue cells are highlighted.", "Admins can edit cells and revalidate without returning to Excel.", "Only fully valid edited batches replace saved requirements."]),
        ("Slide 12 - Review Timetable", ["Timetable grid displays generated sessions.", "Filters support programme, group, staff, room, and day views.", "Schedule versions and explanations make results traceable."]),
        ("Slide 13 - Conflict Repair", ["Hard and soft conflicts are listed.", "Selected sessions show available, warning, and blocked slots.", "Quick fixes suggest clean room/time alternatives.", "The system rechecks conflicts after each move."]),
        ("Slide 14 - Export and Testing", ["CSV/XLSX export matches required system-template columns.", "Tests cover import, validation, solver, conflicts, manual moves, and export.", "Build verification confirms frontend production readiness."]),
        ("Slide 15 - Impact and Future Work", ["Impact: faster scheduling, fewer hidden clashes, export-ready output.", "Future work: authentication, audit history, shared-module modeling, staff/room unavailability, deterministic solving."]),
    ]
    for title, bullets in slide_content:
        doc.add_heading(title, level=2)
        add_bullets(doc, bullets)

    doc.add_heading("4. Demo Flow", level=1)
    add_numbers(
        doc,
        [
            "Open Dashboard and state the current workflow status.",
            "Go to Import Data and upload or show the sample Excel workbook.",
            "Show the import summary and editable preview grid. Mention how errors are highlighted.",
            "Open Generate Timetable and run generation with soft priorities.",
            "Open Review Timetable and show filters, timetable grid, conflicts, and schedule explanations.",
            "Select one conflict or scheduled item and show available/blocked move slots.",
            "Apply a quick fix or manual move if demo data has conflicts.",
            "Open Export and show CSV/XLSX system-template output columns.",
        ],
    )

    doc.add_heading("5. Q&A Preparation", level=1)
    add_table(
        doc,
        ["Possible Question", "Recommended Answer"],
        [
            ["Why use CP-SAT?", "Timetabling is a constraint satisfaction and optimization problem. CP-SAT supports boolean assignment variables, hard constraints, and weighted soft objectives."],
            ["What happens if data has conflicts?", "Invalid import data is rejected before saving. If the timetable is over-constrained, the system can still produce reviewable conflict information instead of silently failing."],
            ["How do you prevent bad uploads?", "The backend validates the entire batch first. If any row fails, no requirement rows are saved until the user fixes and reapplies the data."],
            ["Why can solver results differ?", "CP-SAT can use multiple workers and there can be many equally valid schedules. Deterministic settings can be added as future work."],
            ["What is the main innovation?", "Combining Excel recovery, strict validation, CP-SAT scheduling, conflict repair, and downstream export into one practical workflow."],
        ],
        [3000, 6360],
    )

    doc.add_heading("6. Poster Content", level=1)
    add_table(
        doc,
        ["Poster Section", "Ready-to-Use Content"],
        [
            ["Topic", "Academic Timetable Scheduling System. Team number: To be confirmed. Team: Roy, Xian Yang, Kai Xian, Christie, Ikin, Anastasia. Supervisor: Ms Yang."],
            ["Background", "Academic timetabling requires matching modules, staff, student groups, rooms, delivery modes, teaching weeks, and fixed sessions. Manual planning can be slow and prone to hidden clashes."],
            ["Problem Statement", "The project addresses how to convert timetable requirement spreadsheets into a valid, optimized, reviewable, and export-ready timetable while reducing manual checking effort."],
            ["Methodology / Innovation", "The system uses a React frontend, FastAPI backend, SQLite database, Excel import/export pipeline, and OR-Tools CP-SAT solver. It validates input rows, models hard and soft constraints, and supports conflict review with quick fixes."],
            ["Solution / Impact", "Users can upload Excel requirements, correct errors in-app, generate a timetable, inspect conflicts, repair placements, and export a system-template CSV/XLSX. This reduces repeated manual checks and improves timetable reliability."],
            ["Future Work", "Add authentication, audit history, richer shared-module modeling, staff/room unavailability, deterministic solver settings, and more downstream export templates."],
        ],
        [2100, 7260],
    )

    doc.add_heading("7. Poster Design Suggestion", level=1)
    add_bullets(
        doc,
        [
            "Use a left-to-right pipeline visual: Excel Input -> Validation -> CP-SAT Solver -> Review & Fix -> Export.",
            "Put the problem statement near the top-left and the final impact near the bottom-right.",
            "Include one screenshot each for Import Preview, Review Timetable, and Export format if space allows.",
            "Keep text short; use the longer technical explanation in the presentation instead of overloading the poster.",
        ],
    )
    return doc


def presentation_split():
    doc = make_doc(
        "Presentation Split and Equal Technical Contribution",
        "Balanced six-person speaking plan, technical ownership, demo responsibilities, and backup Q&A map.",
    )

    doc.add_heading("1. Equal Contribution Principle", level=1)
    add_para(
        doc,
        "Each presenter should own a real technical subsystem and speak for roughly the same duration. The split below avoids assigning anyone only introduction or conclusion work; every member explains implementation details that can be defended during Q&A.",
    )

    doc.add_heading("2. Recommended Six-Person Split", level=1)
    add_table(
        doc,
        ["Member", "Primary Technical Area", "Slides", "Speaking Time", "Evidence to Mention"],
        [
            ["Roy", "Problem framing and full-system architecture", "1-4", "4:00", "Workflow pipeline, tech stack, API/database/solver/frontend integration."],
            ["Xian Yang", "Excel import, validation, and database model", "5-6", "3:20", "Input_Template, column normalization, batch validation, SQLite tables, editable preview payload."],
            ["Kai Xian", "CP-SAT optimization and solver performance", "7-9", "4:40", "Decision variables, hard/soft constraints, strict-first solving, relaxed fallback, bucketed constraints."],
            ["Christie", "Frontend workflow and import UX", "10-11", "3:05", "React pages, upload flow, ImportPreviewGrid, readiness panel, soft preference review."],
            ["Ikin", "Timetable review, conflict handling, and quick fixes", "12-13", "3:30", "Review timetable grid, filters, conflict table, manual moves, QuickFixService suggestions."],
            ["Anastasia", "Export, testing, impact, and future work", "14-15", "3:25", "System-template CSV/XLSX mapping, test coverage, known limitations, future improvements."],
        ],
        [1150, 2700, 850, 1050, 3610],
    )

    doc.add_heading("3. Speaker Notes by Member", level=1)
    notes = {
        "Roy": [
            "Start by naming the problem: timetable planning is a constraint-heavy operational task.",
            "Explain the project goal in one sentence: convert Excel requirements into validated, optimized, reviewable, export-ready timetables.",
            "Use the architecture slide to show how the frontend, backend, database, solver, and export pipeline connect.",
        ],
        "Xian Yang": [
            "Explain why clean input matters: a solver cannot produce meaningful results from inconsistent references.",
            "Discuss Input_Template, Remarks_(optional), canonical columns, and all-or-nothing import.",
            "Point out that editable preview rows let users correct errors without returning to Excel.",
        ],
        "Kai Xian": [
            "Define the assignment variable x(session, room, time slot).",
            "Separate hard constraints from soft constraints clearly.",
            "Mention performance: bucketed constraints and strict-first solve reduce unnecessary search pressure.",
            "Explain why fallback mode exists: it keeps over-constrained cases reviewable.",
        ],
        "Christie": [
            "Walk through the UI workflow and connect each page to a user task.",
            "Highlight import preview controls: search, issues-only filter, all-columns toggle, highlighted cells, and validate-and-apply.",
            "Explain how readiness and soft preference review help users prepare before generation.",
        ],
        "Ikin": [
            "Show how the review page changes solver output into actionable decisions.",
            "Explain hard vs soft conflicts from the user's perspective.",
            "Demonstrate or describe manual move and quick-fix suggestions.",
        ],
        "Anastasia": [
            "Explain why export is important: the project must produce a usable downstream output, not just an on-screen timetable.",
            "Name the system-template columns and mapping behavior.",
            "Close with testing, project impact, limitations, and future work.",
        ],
    }
    for member, points in notes.items():
        doc.add_heading(member, level=2)
        add_bullets(doc, points)

    doc.add_heading("4. Demo Responsibility Split", level=1)
    add_table(
        doc,
        ["Demo Step", "Primary", "Backup", "What to Show"],
        [
            ["Project launch and dashboard", "Roy", "Christie", "App opens, workflow navigation, current schedule state."],
            ["Import Excel file", "Xian Yang", "Christie", "Upload box, import summary, preview rows."],
            ["Edit/validate import row", "Christie", "Xian Yang", "Highlighted error cell and validate-and-apply edits."],
            ["Generate timetable", "Kai Xian", "Roy", "Generate Timetable page and solver status result."],
            ["Review generated timetable", "Ikin", "Christie", "Timetable grid, filters, schedule versions, explanations."],
            ["Conflict quick fix", "Ikin", "Kai Xian", "Quick-fix tray, suggested alternatives, recheck behavior."],
            ["Export output", "Anastasia", "Roy", "CSV/XLSX export and system-template column shape."],
        ],
        [2250, 1300, 1300, 4510],
    )

    doc.add_heading("5. Backup Q&A Ownership", level=1)
    add_table(
        doc,
        ["Question Area", "Lead Answer", "Backup"],
        [
            ["Overall purpose and scope", "Roy", "Anastasia"],
            ["Excel input format and validation errors", "Xian Yang", "Christie"],
            ["Database and reference tables", "Xian Yang", "Roy"],
            ["CP-SAT model and constraints", "Kai Xian", "Ikin"],
            ["Solver performance and infeasible cases", "Kai Xian", "Roy"],
            ["Frontend workflow and usability", "Christie", "Ikin"],
            ["Conflict review and quick fixes", "Ikin", "Kai Xian"],
            ["Export format and tests", "Anastasia", "Xian Yang"],
            ["Limitations and future work", "Anastasia", "Roy"],
        ],
        [3400, 2600, 3360],
    )

    doc.add_heading("6. Rehearsal Checklist", level=1)
    add_bullets(
        doc,
        [
            "Keep each member within the planned time window; the full presentation must stay within 20 minutes.",
            "Use consistent terminology: requirements, sessions, rooms, time slots, hard constraints, soft constraints, schedule run, violations.",
            "Do one full demo rehearsal with the same workbook before presentation day.",
            "Prepare a fallback video or screenshots in case live demo setup fails.",
            "Each member should be able to answer at least one Q&A question about their subsystem and one about how it connects to the full pipeline.",
        ],
    )
    return doc


def save_doc(doc, filename):
    path = OUT_DIR / filename
    doc.save(path)
    return path


def main():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    docs = [
        save_doc(technical_breakdown(), "01_Full_Technical_Breakdown.docx"),
        save_doc(presentation_and_poster(), "02_Presentation_and_Poster_Content.docx"),
        save_doc(presentation_split(), "03_Presentation_Split_Equal_Technical_Contribution.docx"),
    ]
    for path in docs:
        print(path)


if __name__ == "__main__":
    main()
