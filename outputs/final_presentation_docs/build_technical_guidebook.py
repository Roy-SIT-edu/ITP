from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "outputs" / "final_presentation_docs" / "01_Full_Technical_Breakdown.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
NAVY = RGBColor(11, 37, 69)
MUTED = RGBColor(89, 89, 89)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
CODE_FILL = "F7F8FA"
BORDER = "C8D1DC"
INK = RGBColor(30, 30, 30)

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_run_font(run, *, name="Calibri", size=None, color=None, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa: list[int], indent_dxa=TABLE_INDENT_DXA):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(0, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths_dxa[min(idx, len(widths_dxa) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "nil")


def style_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "CodeBlock" not in doc.styles:
        code = doc.styles.add_style("CodeBlock", 1)
        code.font.name = "Consolas"
        code._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
        code._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
        code.font.size = Pt(8.5)
        code.font.color.rgb = RGBColor(36, 41, 47)
        code.paragraph_format.space_after = Pt(0)
        code.paragraph_format.line_spacing = 1.05

    set_section_header_footer(doc)


def set_section_header_footer(doc: Document):
    section = doc.sections[0]
    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = ""
    header_p.paragraph_format.space_after = Pt(0)
    left = header_p.add_run("Academic Timetable Scheduling System")
    set_run_font(left, size=9, color=MUTED, bold=True)
    sep = header_p.add_run("  |  Technical Guidebook")
    set_run_font(sep, size=9, color=MUTED)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = ""
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(0)
    footer_p.paragraph_format.space_after = Pt(0)
    run = footer_p.add_run("Team: Roy, Xian Yang, Kai Xian, Christie, Ikin, Anastasia")
    set_run_font(run, size=8.5, color=MUTED)


def add_para(doc, text="", *, style=None, bold_prefix=None, italic=False, color=None, after=None):
    p = doc.add_paragraph(style=style)
    if after is not None:
        p.paragraph_format.space_after = Pt(after)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True, color=color)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_run_font(r2, italic=italic, color=color)
    else:
        r = p.add_run(text)
        set_run_font(r, italic=italic, color=color)
    return p


def add_bullets(doc, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_run_font(run)


def add_numbers(doc, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        run = p.add_run(item)
        set_run_font(run)


def add_table(doc, headers: list[str], rows: list[list[str]], widths_dxa: list[int], *, font_size=9, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, header in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        set_run_font(run, size=font_size, bold=True, color=NAVY)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            run = p.add_run(value)
            set_run_font(run, size=font_size)
    set_table_geometry(table, widths_dxa)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_label_detail_table(doc, rows: list[tuple[str, str]]):
    return add_table(
        doc,
        ["Area", "Technical explanation"],
        [[label, detail] for label, detail in rows],
        [1900, 7460],
        font_size=9,
    )


def add_callout(doc, label: str, text: str):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, LIGHT_GRAY)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    label_run = p.add_run(f"{label}: ")
    set_run_font(label_run, bold=True, color=NAVY)
    run = p.add_run(text)
    set_run_font(run)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_code_block(doc, title: str, code: str, explanation: str | None = None):
    doc.add_heading(title, level=3)
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CODE_FILL)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    p = cell.paragraphs[0]
    p.style = doc.styles["CodeBlock"]
    for index, line in enumerate(code.strip("\n").splitlines()):
        if index:
            p.add_run().add_break()
        run = p.add_run(line.rstrip())
        set_run_font(run, name="Consolas", size=8.5, color=RGBColor(36, 41, 47))
    if explanation:
        add_para(doc, explanation, color=MUTED, italic=True, after=8)
    else:
        doc.add_paragraph().paragraph_format.space_after = Pt(4)


def page_break(doc):
    doc.add_page_break()


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(108)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Academic Timetable Scheduling System")
    set_run_font(r, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(22)
    r = p.add_run("Full Technical Guidebook")
    set_run_font(r, size=18, color=DARK_BLUE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    r = p.add_run("System architecture, APIs, solver design, integrations, code snippets, and implementation rationale")
    set_run_font(r, size=11.5, color=MUTED)

    rows = [
        ("Supervisor", "Ms Yang"),
        ("Team members", "Roy, Xian Yang, Kai Xian, Christie, Ikin, Anastasia"),
        ("Purpose", "Replacement technical breakdown for final presentation, poster, and project assessment"),
        ("Last updated", "9 July 2026"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    for idx, (label, value) in enumerate(rows):
        cells = table.rows[idx].cells
        set_cell_shading(cells[0], LIGHT_BLUE)
        set_cell_shading(cells[1], "FFFFFF")
        cells[0].paragraphs[0].add_run(label).bold = True
        cells[1].paragraphs[0].add_run(value)
        for cell in cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, size=10)
    set_table_geometry(table, [1850, 7510])
    page_break(doc)


def add_information(doc):
    doc.add_heading("Information", level=1)
    add_para(
        doc,
        "This guidebook replaces the current technical breakdown document. It explains the timetable scheduling project as one cohesive system: how Excel requirement rows enter the app, how they become validated database records, how the CP-SAT solver chooses room/time assignments, how conflicts are checked, how manual review works, and how the final timetable is exported.",
    )
    add_para(
        doc,
        "The project is not simply a user interface over a library. The main technical contribution is translating real-world academic timetabling into a reliable software pipeline: reference-data management, validation, candidate filtering, constraint modelling, solver fallback behaviour, post-generation auditing, and a review/export workflow that non-technical users can operate.",
    )
    add_callout(
        doc,
        "Core technical idea",
        "Each requirement session is converted into compatible room/time candidates. The solver chooses exactly one candidate per session, blocks hard clashes where possible, scores soft preferences, and then the system audits the result so users can inspect and fix remaining conflicts before export.",
    )
    add_label_detail_table(
        doc,
        [
            (
                "Frontend",
                "React + TypeScript + Vite single-page app. It gives users upload, database admin, soft-constraint ranking, generation, timetable review, conflict fixing, and export screens.",
            ),
            (
                "Backend API",
                "FastAPI service with routers for upload, validation, reference data, database management, soft constraints, schedule generation, and export.",
            ),
            (
                "Persistence",
                "SQLAlchemy models backed by local SQLite. It stores master reference tables, imported requirement sessions, generated schedule runs, assigned sessions, and stored violations.",
            ),
            (
                "Solver",
                "Google OR-Tools CP-SAT. It builds boolean decision variables for feasible session/slot/room assignments and minimizes weighted penalties.",
            ),
            (
                "Output",
                "CSV/XLSX system-template export, blocked until hard conflicts are resolved.",
            ),
        ]
    )


def add_contents(doc):
    doc.add_heading("Guidebook Contents", level=1)
    sections = [
        "Overall system architecture",
        "Runtime stack and repository structure",
        "Frontend-backend API integration",
        "Backend API catalogue",
        "Database model and stored entities",
        "Excel import and editable preview pipeline",
        "Requirement validation and reference resolution",
        "Schedule generation orchestration",
        "CP-SAT model construction",
        "Hard constraints, soft constraints, and ranking weights",
        "Solver execution, relaxed solving, and fallback behaviour",
        "Result parsing, persistence, and post-generation audit",
        "Timetable review, manual move, and quick-fix integrations",
        "Export pipeline and system-template mapping",
        "Frontend workflow implementation",
        "Testing, verification, limitations, and future work",
        "Appendix: file map and API summary",
    ]
    add_numbers(doc, sections)
    page_break(doc)


def add_architecture(doc):
    doc.add_heading("Overall System Architecture", level=1)
    add_para(
        doc,
        "The system follows a layered architecture. Each layer has one job, and the integration between layers is explicit: the frontend never talks directly to the database or solver, the solver never parses Excel, and export only reads persisted schedule results that have already passed the conflict workflow.",
    )
    add_callout(
        doc,
        "Architecture flow",
        "Excel workbook -> Upload API -> ImportService -> RequirementInputService -> SQLite tables -> ValidationService -> ScheduleService -> CP-SAT Solver -> ScheduledSession rows -> ConstraintService -> Review UI -> ExportService -> CSV/XLSX.",
    )
    add_table(
        doc,
        ["Step", "Component", "Purpose", "Integration point"],
        [
            ["1", "Excel input workbook", "Supplies timetable requirements, staff IDs, class sizes, fixed times, preferred days, and remarks.", "Uploaded from React as multipart form data."],
            ["2", "Upload router + ImportService", "Reads workbook sheets, normalizes column aliases, produces editable preview rows, and blocks invalid batches.", "Calls pandas/openpyxl and RequirementInputService."],
            ["3", "RequirementInputService", "Resolves external codes to database IDs, validates fields, creates Session and SessionStaff records.", "Uses SQLAlchemy models and shared compatibility helpers."],
            ["4", "ValidationService", "Checks saved requirements before generation and summarizes readiness/conflicts.", "Feeds dashboard/readiness UI."],
            ["5", "ScheduleService", "Creates a ScheduleRun, loads sessions/slots/rooms/soft weights, invokes the solver, persists assignments.", "Calls CpSatTimetableSolver and ConstraintService."],
            ["6", "CP-SAT solver layer", "Builds candidate variables, hard no-overlap rules, soft penalties, objective, fallback solving.", "Uses OR-Tools and shared scheduling rules."],
            ["7", "ConstraintService", "Audits generated or manually edited schedules for hard violations and soft warnings.", "Stores ConstraintViolation rows and updates schedule status."],
            ["8", "Review and export UI", "Lets users inspect runs, move sessions, request quick fixes, and download clean output.", "Calls schedule, violation, quick-fix, and export APIs."],
        ],
        [650, 1900, 3550, 3260],
        font_size=8.5,
    )


def add_stack(doc):
    doc.add_heading("Runtime Stack and Repository Structure", level=1)
    add_table(
        doc,
        ["Layer", "Technology", "Reason for use", "Important paths"],
        [
            ["Frontend", "React, TypeScript, Vite", "Fast local SPA with typed API contracts and componentized review/edit workflows.", "frontend/src/App.tsx, frontend/src/api/client.ts, frontend/src/pages"],
            ["API", "FastAPI", "Clear router structure, Pydantic validation, async-friendly file upload, simple OpenAPI surface.", "backend/app/main.py, backend/app/routes"],
            ["Database", "SQLAlchemy + SQLite", "Relational storage for reference data, sessions, schedule runs, assignments, and violations without external DB setup.", "backend/app/models, backend/app/database.py"],
            ["Excel processing", "pandas, openpyxl", "Robust reading/writing of workbook templates and system-template export.", "ImportService, ExportService"],
            ["Solver", "Google OR-Tools CP-SAT", "Models assignment decisions and constraints as boolean/integer optimization, which fits timetabling.", "backend/app/solver"],
            ["Verification", "pytest, ruff, frontend build/lint", "Backend logic tests and frontend static/build checks for regression control.", "backend/tests, frontend package scripts"],
        ],
        [1300, 1700, 3900, 2460],
        font_size=8.5,
    )
    add_code_block(
        doc,
        "Backend application wiring",
        """
app = FastAPI(title="Timetable Scheduling API", version="0.1.0", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
    allow_origin_regex=r"http://(localhost|127\\.0\\.0\\.1):\\d+",
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

app.include_router(upload_routes.router)
app.include_router(validation_routes.router)
app.include_router(data_routes.router)
app.include_router(database_routes.router)
app.include_router(soft_constraint_routes.router)
app.include_router(schedule_routes.router)
app.include_router(export_routes.router)
""",
        "FastAPI is the integration hub. The routers keep each workflow separate while CORS allows the Vite development frontend to call the backend from changing local ports.",
    )


def add_api_integration(doc):
    doc.add_heading("Frontend-Backend API Integration", level=1)
    add_para(
        doc,
        "The frontend uses one central client so every page gets consistent URL construction, JSON parsing, multipart upload support, export URLs, and error handling. This is important because FastAPI can return simple strings, Pydantic validation arrays, or structured row-level import errors.",
    )
    add_code_block(
        doc,
        "Central frontend request wrapper",
        """
export const API_BASE = import.meta.env.VITE_API_URL ?? "";

async function request<T>(path: string, options?: RequestInit): Promise<T> {
  const response = await fetch(`${API_BASE}${path}`, options);
  const contentType = response.headers.get("content-type") ?? "";
  const payload = contentType.includes("application/json")
    ? await response.json()
    : await response.text();

  if (!response.ok) {
    const detail = typeof payload === "object" && payload && "detail" in payload ? payload.detail : payload;
    const message = typeof detail === "string" ? detail : `Request failed with status ${response.status}`;
    throw new ApiError(response.status, message, detail);
  }
  return payload as T;
}
""",
        "Purpose: avoid duplicating fetch/error code across Upload, Database, Soft Constraints, Review, Dashboard, and Export pages. Integration: every exported client function maps directly to a backend route.",
    )
    add_table(
        doc,
        ["API group", "Key endpoints", "Purpose", "Main service/model integration"],
        [
            ["Upload", "POST /api/upload/input-template; /preview; /edited", "Import Excel files, preview row issues, and apply user-corrected rows.", "ImportService, RequirementInputService, Session, SessionStaff"],
            ["Data", "GET /api/dashboard; /sessions; /rooms; /timeslots; /availability; /constraint-insights", "Feed workflow screens, readiness panels, manual forms, and availability summaries.", "Serializers, ValidationService, ScheduleRun, ConstraintViolation"],
            ["Database", "GET/POST/PUT/DELETE /api/database/{type}; upload/example/current", "Manage reference data used by validation and solving.", "DatabaseService and master models"],
            ["Validation", "GET /api/validation/latest", "Show whether uploaded sessions are ready to generate.", "ValidationService"],
            ["Soft constraints", "GET/PUT /api/soft-constraints", "Let users rank and disable soft preferences before solving.", "SoftConstraintPriorityService"],
            ["Schedules", "POST /api/schedules/generate; GET latest/runs/compare; PUT move; POST suggest-fixes/recheck", "Generate, inspect, compare, move, fix, and re-audit timetables.", "ScheduleService, CpSatTimetableSolver, ConstraintService, QuickFixService"],
            ["Export", "GET /api/export/{run}/csv; /xlsx", "Download final system-template files only after hard conflicts are cleared.", "ExportService, ScheduleRun"],
        ],
        [1350, 2650, 3000, 2360],
        font_size=8.3,
    )


def add_database_model(doc):
    doc.add_heading("Database Model and Stored Entities", level=1)
    add_para(
        doc,
        "The database separates stable reference data from generated results. That separation matters because users can change rooms/staff/modules without rewriting solver code, while each generated timetable remains traceable through a ScheduleRun and its ScheduledSession rows.",
    )
    add_table(
        doc,
        ["Entity/table", "What it stores", "Why it exists / integration"],
        [
            ["programmes, modules, student_groups, staff, rooms, time_slots", "Master data used by uploaded rows and solver candidates.", "Prevents free-text scheduling decisions. Validation resolves codes into IDs before generation."],
            ["sessions", "Imported or manually entered timetable requirements: module, group, staff, delivery mode, room needs, week pattern, fixed/preferred/avoid rules.", "Main requirement table consumed by ValidationService and CpSatTimetableSolver."],
            ["session_staff", "Primary and co-teaching staff links for a session.", "Lets staff clash checks include Staff 1 through Staff 4 instead of only one teacher."],
            ["lab_requirements", "Built-in fixed lab bookings and related metadata.", "Synced into sessions as hard constraints before generation."],
            ["soft_constraint_priorities", "Soft constraint rank, weight, and active/inactive state.", "Converts user ranking into solver objective weights."],
            ["schedule_runs", "One row per generation attempt: status, solver status, hard violation count, soft score, message.", "Lets the app compare runs and keep review/export tied to a specific result."],
            ["scheduled_sessions", "The chosen room/time/session assignments for a schedule run.", "Review UI, conflict checker, explanations, and export all read this table."],
            ["constraint_violations", "Stored hard/soft issues after a solve or manual move.", "Makes conflicts visible and persistent for review and export gating."],
        ],
        [2100, 3600, 3660],
        font_size=8.7,
    )
    add_code_block(
        doc,
        "Core schedule result models",
        """
class ScheduleRun(Base):
    __tablename__ = "schedule_runs"
    id = Column(Integer, primary_key=True, index=True)
    status = Column(String, nullable=False, default="PENDING")
    solver_status = Column(String, nullable=True)
    hard_violation_count = Column(Integer, nullable=False, default=0)
    soft_score = Column(Integer, nullable=False, default=0)
    message = Column(String, nullable=True)

class ScheduledSession(Base):
    __tablename__ = "scheduled_sessions"
    schedule_run_id = Column(Integer, ForeignKey("schedule_runs.id"), nullable=False)
    session_id = Column(Integer, ForeignKey("sessions.id"), nullable=False)
    room_id = Column(Integer, ForeignKey("rooms.id"), nullable=False)
    time_slot_id = Column(Integer, ForeignKey("time_slots.id"), nullable=False)
    day = Column(String, nullable=False)
    start_time = Column(String, nullable=False)
    end_time = Column(String, nullable=False)
""",
        "Purpose: keep generated timetables auditable. Integration: ScheduleService writes these rows, ConstraintService checks them, TimetableReviewPage reads them, and ExportService converts them into the required output format.",
    )


def add_import_pipeline(doc):
    doc.add_heading("Excel Import and Editable Preview Pipeline", level=1)
    add_para(
        doc,
        "The import layer accepts real Excel workbooks rather than requiring perfectly formatted JSON. It normalizes column names, supports a documented Input_Template sheet plus optional remarks, and provides an editable preview so users can fix bad rows without leaving the app.",
    )
    add_label_detail_table(
        doc,
        [
            ("Purpose", "Convert inconsistent workbook input into a clean list of RequirementUploadRow objects."),
            ("Reason", "Real timetable files often vary in header spelling and may contain blank rows or optional sheets."),
            ("Integration", "Upload routes call ImportService, which calls RequirementInputService. Valid imports clear previous sessions and schedule runs to avoid mixing old and new requirements."),
        ]
    )
    add_code_block(
        doc,
        "Column alias normalization",
        """
ALIAS_LOOKUP = {
    _normalise_column_name(alias): canonical
    for canonical, aliases in CANONICAL_COLUMNS.items()
    for alias in [canonical, *aliases]
}

REQUIRED_TEMPLATE_COLUMNS = [
    "Requirement ID",
    "Programme",
    "Year",
    "Module Code",
    "Class Type",
    "Session Count",
    "Duration Hours",
    "Sessions Per Week",
    "Delivery Mode",
    "Venue Type Required",
    "Exact Class Size",
    "Staff 1 ID",
]
""",
        "Why this matters: users can upload workbooks with headers such as 'staff id', 'sis staff id', or 'Staff 1 ID', and the backend still maps them to one canonical field.",
    )
    add_code_block(
        doc,
        "Workbook sheet selection",
        """
def _read_workbook(self, workbook_source) -> PreparedWorkbook:
    with pd.ExcelFile(workbook_source) as xls:
        if "Input_Template" in xls.sheet_names:
            required = pd.read_excel(xls, sheet_name="Input_Template").dropna(how="all")
            optional = (
                pd.read_excel(xls, sheet_name="Remarks_(optional)").dropna(how="all")
                if "Remarks_(optional)" in xls.sheet_names
                else pd.DataFrame()
            )
            return self._prepare_two_tab_workbook(required, optional)

        sheet_name = self._choose_sheet(xls.sheet_names)
        frame = pd.read_excel(xls, sheet_name=sheet_name).dropna(how="all")
        prepared = self._prepare_frame(frame, require_documented_shape=False)
        return PreparedWorkbook(frame=prepared, errors=[], rows_read=int(len(prepared.index)), columns=list(prepared.columns))
""",
        "Purpose: support both the official two-sheet template and simpler one-sheet workbooks. Integration: the returned PreparedWorkbook is used for preview and final import.",
    )
    add_code_block(
        doc,
        "Edited preview rows become the final import",
        """
def import_preview_rows(self, db: DbSession, rows: list[dict]) -> dict:
    upload_rows = []
    for index, row in enumerate(rows):
        values = row.get("values") or {}
        source_filename = clean_text(row.get("source_file")) or "Edited import"
        source_row_no = self._source_row_no(row.get("source_row_no"), index)
        upload_rows.append(RequirementUploadRow(row=values, source_filename=source_filename, source_row_no=source_row_no))

    session_data, validation_errors = RequirementInputService().validate_upload_rows(db, upload_rows)
    if validation_errors:
        db.rollback()
        return {"rows_imported": 0, "errors": validation_errors, "preview_rows": preview_rows}

    self._clear_sessions_and_schedules(db)
    for data in session_data:
        db.add(service.session_from_data(data))
    db.commit()
""",
        "Reason: import is all-or-nothing. A bad edited row does not partially replace the existing requirement set.",
    )


def add_requirement_validation(doc):
    doc.add_heading("Requirement Validation and Reference Resolution", level=1)
    add_para(
        doc,
        "RequirementInputService is the bridge between human spreadsheet data and relational solver-ready data. It rejects duplicate requirement IDs, verifies required fields, resolves programme/module/group/staff references, supports co-teachers, checks fixed time slots, and confirms that at least one feasible room exists.",
    )
    add_code_block(
        doc,
        "Upload validation loop",
        """
def validate_upload_rows(self, db: DbSession, rows: list[RequirementUploadRow]):
    errors = []
    session_data = []
    seen_requirement_ids = {}

    for item in rows:
        requirement_id = clean_text(self._value(item.row, "Requirement ID"))
        if requirement_id:
            key = requirement_id.lower()
            if key in seen_requirement_ids:
                errors.append(self._issue(item.source_row_no, "Requirement ID", "Duplicate requirement_id in upload."))
            else:
                seen_requirement_ids[key] = item.source_row_no

        data, row_errors = self._build_session_data(
            db, item.row, item.source_filename, item.source_row_no,
            check_existing_duplicate=False,
            allow_reference_upsert=True,
        )
        errors.extend(row_errors)
        if not row_errors:
            session_data.append(data)

    return session_data, errors
""",
        "Purpose: collect every row issue into a user-readable preview instead of failing at the first error. Integration: ImportPreviewGrid displays these row/field errors.",
    )
    add_code_block(
        doc,
        "Co-teacher support",
        """
def _staff_assignments(self, db: DbSession, row: Mapping[str, Any], source_row_no: int, errors: list[dict]):
    assignments = []
    seen_staff_ids = {}
    for staff_order in range(1, 5):
        id_field = f"Staff {staff_order} ID"
        staff_id = clean_text(self._value(row, id_field))
        if staff_order == 1 and not staff_id:
            errors.append(self._issue(source_row_no, id_field, "Staff 1 ID is required."))
            continue
        if not staff_id:
            continue
        staff = self._lookup_staff_by_id(db, staff_id, source_row_no, id_field, errors)
        if staff and staff.id not in seen_staff_ids:
            assignments.append({"staff": staff, "staff_order": staff_order, "is_primary": staff_order == 1})
    return assignments
""",
        "Reason: timetable clashes must consider every teacher assigned to a class, not only the primary staff member.",
    )
    add_code_block(
        doc,
        "Feasible-room check reuses solver compatibility",
        """
def _has_feasible_room(self, db: DbSession, delivery_mode, campus_mode, venue_type, exact_class_size) -> bool:
    probe = SimpleNamespace(
        delivery_mode=delivery_mode,
        campus_mode=campus_mode,
        venue_type_required=venue_type,
        exact_class_size=exact_class_size,
    )
    for room in db.query(Room).all():
        if not self._campus_room_compatible(campus_mode, room):
            continue
        if delivery_room_compatible(probe, room) and venue_room_compatible(probe, room) and room_capacity_fits(probe, room):
            return True
    return False
""",
        "Purpose: validation and solving agree on room feasibility because they share the same compatibility helpers.",
    )


def add_schedule_service(doc):
    doc.add_heading("Schedule Generation Orchestration", level=1)
    add_para(
        doc,
        "ScheduleService is the backend workflow coordinator. It is intentionally separate from the solver: the service handles database state, lab requirement synchronization, soft-priority lookup, persistence, and post-generation checks, while the solver only returns candidate assignments and a score.",
    )
    add_code_block(
        doc,
        "Generation service sequence",
        """
def generate(self, db: DbSession, timeout=DEFAULT_GENERATION_TIMEOUT_SECONDS, fast_mode=False) -> dict:
    active_lab_requirement_ids = self.lab_requirement_service.sync_active_to_sessions(db)
    db.commit()

    run = ScheduleRun(status="RUNNING", message="Solver started")
    db.add(run)
    db.commit()
    run_id = run.id

    sessions = [
        item for item in db.query(Session).order_by(Session.id).all()
        if not item.is_lab_requirement or item.requirement_id in active_lab_requirement_ids
    ]
    time_slots = db.query(TimeSlot).order_by(TimeSlot.day, TimeSlot.start_time).all()
    rooms = db.query(Room).order_by(Room.room_code).all()
    soft_weights = self.priority_service.weights(db)

    result = self.solver.solve(sessions, time_slots, rooms, soft_constraint_weights=soft_weights, max_seconds=timeout, fast_mode=fast_mode)
""",
        "Integration: the service creates the ScheduleRun first so every assignment and violation can be tied to a generation attempt.",
    )
    add_numbers(
        doc,
        [
            "Sync active lab requirements into solver-visible sessions.",
            "Create a RUNNING ScheduleRun row.",
            "Load sessions, time slots, rooms, and soft-constraint weights.",
            "Call CpSatTimetableSolver.",
            "If solver fails, mark the run FAILED and return the solver message.",
            "If solver succeeds, persist ScheduledSession rows.",
            "Run ConstraintService.check_and_store and update hard_violation_count, soft_score, and status.",
        ],
    )


def add_solver_model(doc):
    doc.add_heading("CP-SAT Model Construction", level=1)
    add_para(
        doc,
        "The solver model is a decision-variable grid. For each requirement session, the builder creates a boolean variable only for room/time combinations that pass shared candidate rules. This keeps invalid assignments out of the model before optimization even begins.",
    )
    add_label_detail_table(
        doc,
        [
            ("Decision variable", "x_session_slot_room is true when the solver chooses that exact session, time slot, and room."),
            ("Candidate filtering", "candidate_slot_allowed checks duration, week pattern, fixed time, and hard avoid days. candidate_room_allowed checks required room codes, capacity, delivery mode, and venue type."),
            ("Exact-once rule", "Every schedulable session must have exactly one selected candidate."),
            ("Objective", "Minimize weighted soft penalties plus heavy penalties only in relaxed/fallback modes."),
        ]
    )
    add_code_block(
        doc,
        "Candidate variables and exact-once constraint",
        """
for session in sessions:
    session_vars = []
    for slot in time_slots:
        if not candidate_slot_allowed(session, slot):
            continue
        for room in rooms:
            if not candidate_room_allowed(session, room):
                continue
            key = (session.id, slot.id, room.id)
            variable = model.NewBoolVar(f"x_{session.id}_{slot.id}_{room.id}")
            variables[key] = variable
            assignments.append({"session": session, "time_slot": slot, "room": room, "variable": variable})
            session_vars.append(variable)

    if not session_vars:
        no_candidate_reasons.append(f"No feasible time slot and room combination is available for {label}.")
        continue

    model.Add(sum(session_vars) == 1)
""",
        "Purpose: make the optimization problem explicit and auditable. If a session has no possible candidates, the backend can report that directly instead of producing a mysterious solver failure.",
    )


def add_constraints(doc):
    doc.add_heading("Hard Constraints, Soft Constraints, and Ranking Weights", level=1)
    add_para(
        doc,
        "The project separates constraints into hard constraints and soft preferences. Hard constraints protect timetable feasibility, while soft constraints express preferences that can be ranked by users before generation.",
    )
    add_table(
        doc,
        ["Type", "Examples", "Where enforced", "Purpose"],
        [
            ["Hard", "Room double booking, staff double booking, student group double booking, room capacity, delivery/room mismatch, fixed day/time.", "Candidate rules, CP-SAT no-overlap constraints, ConstraintService audit, manual move validation.", "Prevent or expose timetable conflicts that must be resolved before export."],
            ["Soft", "Preferred days, avoid days when not hard priority, tutor idle gaps, long consecutive student days, online/F2F adjacent switch, short campus day, online not Mon/Tue.", "CP-SAT objective where feasible, ConstraintService weighted warnings after generation.", "Improve timetable quality without blocking every possible schedule."],
        ],
        [1000, 3600, 2800, 1960],
        font_size=8.5,
    )
    add_code_block(
        doc,
        "Shared candidate rules",
        """
def candidate_slot_allowed(session: Session, slot: TimeSlot) -> bool:
    if session.duration_minutes and slot.duration_minutes != session.duration_minutes:
        return False
    if normalize_token(session.scheduling_type) == "fixed":
        if session.fixed_day and slot.day != session.fixed_day:
            return False
        if session.fixed_start_time and slot.start_time != session.fixed_start_time:
            return False
        if session.fixed_end_time and slot.end_time != session.fixed_end_time:
            return False
    if normalize_token(session.priority) == "hard" and slot.day in parse_day_list(session.avoid_days):
        return False
    return True

def candidate_room_allowed(session: Session, room: Room) -> bool:
    required_codes = required_room_codes(session)
    if required_codes:
        return room.room_code.lower() in {code.lower() for code in required_codes} and delivery_room_compatible(session, room)
    return room_capacity_fits(session, room) and delivery_room_compatible(session, room) and venue_room_compatible(session, room)
""",
        "Reason: validation, solving, and quick-fix suggestions use the same rules, which reduces drift between 'valid input' and 'schedulable input'.",
    )
    add_code_block(
        doc,
        "No-overlap hard rule with relaxed penalty option",
        """
def _add_resource_bucket_rule(self, model, items, penalties, label, relax_hard_conflicts) -> None:
    unique_items = list({id(item["variable"]): item for item in items}.values())
    if len(unique_items) <= 1:
        return
    lab_variables = [item["variable"] for item in unique_items if item["session"].is_lab_requirement]
    non_lab_variables = [item["variable"] for item in unique_items if not item["session"].is_lab_requirement]
    if relax_hard_conflicts:
        self._add_bucket_excess_penalty(model, non_lab_variables, penalties, label)
        for non_lab_variable in non_lab_variables:
            for lab_variable in lab_variables:
                self._add_pair_excess_penalty(model, non_lab_variable, lab_variable, penalties, label)
    else:
        if len(non_lab_variables) > 1:
            model.Add(sum(non_lab_variables) <= 1)
        for non_lab_variable in non_lab_variables:
            for lab_variable in lab_variables:
                model.Add(non_lab_variable + lab_variable <= 1)
""",
        "Purpose: strict mode prevents overlap. Relaxed mode still creates a reviewable timetable by converting unavoidable clashes into large penalties that ConstraintService later reports.",
    )
    add_code_block(
        doc,
        "User ranking becomes solver weight",
        """
def update_priorities(self, db: DbSession, ordered_codes: list[str], active_codes: list[str] | None = None):
    active_order = [code for code in cleaned if code in active_set]
    inactive_order = [code for code in cleaned if code not in active_set]
    total = len(active_order)
    for rank, code in enumerate(active_order, start=1):
        row.rank = rank
        row.weight = self.weight_for_rank(rank, total)
        row.is_active = True
    for rank, code in enumerate(inactive_order, start=total + 1):
        row.rank = rank
        row.weight = 0
        row.is_active = False

@staticmethod
def weight_for_rank(rank: int, total: int) -> int:
    return max(1, total - rank + 1) * 5
""",
        "Integration: SoftConstraintsPage sends ordered/active codes; SoftConstraintPriorityService converts them into integer weights; the model builder uses those weights in the objective.",
    )


def add_solver_execution(doc):
    doc.add_heading("Solver Execution, Relaxed Solving, and Fallback Behaviour", level=1)
    add_para(
        doc,
        "CpSatTimetableSolver is a facade around OR-Tools. It first attempts a strict CP-SAT solve. If there is a known fixed hard clash, a timeout, or strict infeasibility, the code keeps the review workflow usable by producing a reviewable timetable through relaxed or greedy fallback paths, then relies on ConstraintService to expose conflicts.",
    )
    add_code_block(
        doc,
        "Strict solve then fallback strategy",
        """
built = self.model_builder.build(sessions, time_slots, rooms, soft_constraint_weights)
if built.no_candidate_reasons:
    return {"solver_status": "INFEASIBLE", "assignments": [], "message": " ".join(built.no_candidate_reasons)}

result = self._solve_built_model(built, max_seconds, fast_mode)
if result["solver_status"] in {"OPTIMAL", "FEASIBLE"}:
    return result

if result["solver_status"] == "UNKNOWN":
    return self._greedy_fallback(sessions, time_slots, rooms, soft_constraint_weights, "Solver timed out; generated a reviewable timetable with conflict checks.")

if result["solver_status"] == "INFEASIBLE":
    relaxed = self.model_builder.build(sessions, time_slots, rooms, soft_constraint_weights, relax_hard_conflicts=True)
    relaxed_result = self._solve_built_model(relaxed, max_seconds, fast_mode=True)
    if relaxed_result["solver_status"] in {"OPTIMAL", "FEASIBLE"}:
        return relaxed_result
    return self._greedy_fallback(sessions, time_slots, rooms, soft_constraint_weights, "Solver timed out; generated a reviewable timetable with conflict checks.")
""",
        "Reason: a hard 'no timetable' response is accurate but not always useful. A reviewable result lets the team demonstrate conflict detection and manual fixing even when input data is over-constrained.",
    )
    add_code_block(
        doc,
        "CP-SAT runtime parameters",
        """
solver = cp_model.CpSolver()
if max_seconds > 0:
    solver.parameters.max_time_in_seconds = max_seconds
if fast_mode:
    solver.parameters.stop_after_first_solution = True
solver.parameters.num_search_workers = 8
status = solver.Solve(built.model)
""",
        "Purpose: bound generation time and use parallel search workers. Fast mode is useful when the goal is to find any reviewable feasible assignment quickly.",
    )
    add_code_block(
        doc,
        "Selected variable parser",
        """
def parse(self, solver: cp_model.CpSolver, assignments: list[dict]) -> list[dict]:
    results = []
    for assignment in assignments:
        if solver.BooleanValue(assignment["variable"]):
            session = assignment["session"]
            slot = assignment["time_slot"]
            room = assignment["room"]
            results.append({
                "session_id": session.id,
                "room_id": room.id,
                "time_slot_id": slot.id,
                "staff_id": session.staff_id,
                "day": slot.day,
                "start_time": slot.start_time,
                "end_time": slot.end_time,
                "week_pattern": slot.week_pattern,
            })
    return results
""",
        "Integration: ResultParser converts solver internals into plain dictionaries that ScheduleService can persist as ScheduledSession rows.",
    )


def add_constraint_audit(doc):
    doc.add_heading("Post-Generation Constraint Audit", level=1)
    add_para(
        doc,
        "ConstraintService is the independent audit layer. Even if CP-SAT prevents a rule during strict solving, the audit still checks the persisted result. This is essential because relaxed solving, greedy fallback, and manual moves can create schedules that are intentionally reviewable but not yet export-ready.",
    )
    add_code_block(
        doc,
        "Store violations and weighted soft score",
        """
def check_and_store(self, db: DbSession, schedule_run_id: int, soft_constraint_weights=None) -> dict:
    db.query(ConstraintViolation).filter_by(schedule_run_id=schedule_run_id).delete()
    violations = self.check_schedule(db, schedule_run_id)
    for violation in violations:
        db.add(ConstraintViolation(
            schedule_run_id=schedule_run_id,
            constraint_code=violation["constraint_code"],
            severity=violation["severity"],
            message=violation["message"],
            affected_session_ids=",".join(str(item) for item in violation["affected_session_ids"]),
        ))
    hard_count = sum(1 for item in violations if item["severity"] == "HARD")
    weighted_soft_score = sum(weights.get(item["constraint_code"], DEFAULT_SOFT_CONSTRAINT_WEIGHTS.get(item["constraint_code"], 1))
                              for item in violations if item["severity"] == "SOFT")
    return {"violations": violations, "hard_violation_count": hard_count, "weighted_soft_score": weighted_soft_score}
""",
        "Purpose: make conflict state persistent and visible. Integration: review pages load stored violations, export blocks when hard_violation_count is above zero.",
    )
    add_table(
        doc,
        ["Audit area", "Checks performed"],
        [
            ["Room/resource clashes", "Room double booking and lab/non-lab room conflicts with week and time overlap awareness."],
            ["People clashes", "Staff double booking across primary and co-teachers."],
            ["Student clashes", "Student group double booking, including additional required group codes."],
            ["Quality hard checks", "Room capacity mismatch, delivery room mismatch, invalid fixed time placement."],
            ["Soft checks", "Tutor idle gap, short campus day, long consecutive student day, online/F2F adjacent switch, online not Monday/Tuesday."],
        ],
        [2200, 7160],
        font_size=9,
    )


def add_review_and_export(doc):
    doc.add_heading("Timetable Review, Manual Move, Quick Fix, and Export Integrations", level=1)
    add_para(
        doc,
        "After generation, the system behaves like an operational review tool. Users can inspect runs, compare quality, move a class manually, request suggested fixes, re-check constraints, and export only when hard conflicts are resolved.",
    )
    add_code_block(
        doc,
        "Manual move is validated before commit",
        """
item.room_id = room.id
item.time_slot_id = slot.id
item.day = slot.day
item.start_time = slot.start_time
item.end_time = slot.end_time
item.week_pattern = slot.week_pattern

if item.session and item.session.scheduling_type.strip().lower() == "fixed":
    item.session.scheduling_type = "Standard"
    item.session.fixed_day = None
    item.session.fixed_start_time = None
    item.session.fixed_end_time = None

preview_violations = ConstraintService().check_schedule(db, schedule_run_id)
blocking_violations = _hard_violations_for_session(preview_violations, session_id)
if blocking_violations:
    db.rollback()
    raise HTTPException(status_code=409, detail={"message": _manual_move_blocked_message(blocking_violations)})
""",
        "Purpose: manual edits are allowed, but the backend refuses a move that creates a hard conflict for that session. Integration: TimetableReviewPage calls moveScheduledSession and displays conflict messages from the API.",
    )
    add_table(
        doc,
        ["Review capability", "Backend integration", "Why it matters"],
        [
            ["Latest schedule and run history", "GET /api/schedules/latest and /api/schedules", "Users can review the current result and compare recent generation attempts."],
            ["Manual move", "PUT /api/schedules/{run}/sessions/{session}", "Allows human correction while still checking hard rules server-side."],
            ["Quick fix suggestions", "POST /api/schedules/{run}/suggest-fixes", "Ranks alternative slots/rooms that are compatible with the target session."],
            ["Recheck", "POST /api/schedules/{run}/recheck", "Rebuilds violation rows after changes."],
            ["Explanations", "GET /api/schedules/{run}/explanations", "Gives review UI human-readable reasons for scheduled placements."],
        ],
        [2200, 3000, 4160],
        font_size=8.7,
    )
    add_code_block(
        doc,
        "Export gate blocks unresolved hard conflicts",
        """
def _ensure_run(db: DbSession, schedule_run_id: int) -> ScheduleRun:
    run = db.query(ScheduleRun).filter_by(id=schedule_run_id).first()
    if not run:
        raise HTTPException(status_code=404, detail={"message": "Schedule run not found"})
    if int(run.hard_violation_count or 0) > 0:
        raise HTTPException(
            status_code=409,
            detail={
                "message": "You must resolve all Hard Conflicts before you can export your timetable.",
                "hard_conflicts": run.hard_violation_count,
            },
        )
    return run
""",
        "Reason: exported timetables should be safe to submit/use. Soft warnings can remain as quality trade-offs, but hard conflicts must be resolved.",
    )
    add_code_block(
        doc,
        "System-template export columns",
        """
SYSTEM_TEMPLATE_COLUMNS = [
    "Module", "Class Type", "Template", "Group", "Day", "Start", "End",
    "Class Size", "Sector", "RoomGrouping", "Room1", "Room2",
    "StaffGrouping", "Staff1", "Staff2", "Tri Week", "Recording Mode", "Remark",
]
""",
        "Integration: ExportService maps ScheduledSession rows into this exact column layout, converts days to Mon/Tue labels, formats times as HHMM, maps teaching weeks, staff, rooms, class groups, recording mode, and combined-module remarks.",
    )


def add_frontend_workflow(doc):
    doc.add_heading("Frontend Workflow Implementation", level=1)
    add_para(
        doc,
        "The frontend is organized around the real user workflow: dashboard, upload, database maintenance, soft-constraint ranking, timetable review, export, and settings. Hash routing keeps the app simple while still supporting direct links to major screens.",
    )
    add_code_block(
        doc,
        "Hash route map",
        """
const routeMap = {
  dashboard: DashboardPage,
  upload: UploadPage,
  "database-rooms": () => <DatabasePage dataType="rooms" />,
  "database-staff": () => <DatabasePage dataType="staff" />,
  "database-programmes": () => <DatabasePage dataType="programmes" />,
  "database-modules": () => <DatabasePage dataType="modules" />,
  "database-student-groups": () => <DatabasePage dataType="student-groups" />,
  "database-lab-requirements": () => <DatabasePage dataType="lab-requirements" />,
  "soft-constraints": SoftConstraintsPage,
  review: TimetableReviewPage,
  export: ExportPage,
  settings: SettingsPage,
};
""",
        "Purpose: each workflow stage has a dedicated page, but the app remains a single-page application.",
    )
    add_code_block(
        doc,
        "Editable import preview grid",
        """
const visibleColumns = useMemo(() => {
  if (showAllColumns) return allColumns;
  const defaults = DEFAULT_COLUMNS.filter((column) => allColumns.includes(column));
  const issueColumns = issueFields.filter((field) => allColumns.includes(field) && !defaults.includes(field));
  return [...defaults, ...issueColumns];
}, [allColumns, issueFields, showAllColumns]);

const updateCell = (rowId: string, column: string, value: string) => {
  setDraftRows((rows) =>
    rows.map((row) =>
      row.row_id === rowId ? { ...row, values: { ...row.values, [column]: value.trim() === "" ? null : value } } : row,
    ),
  );
  setDirty(true);
};
""",
        "Reason: the upload workflow is not just file input. Users can filter to problem rows, edit cells, and resubmit corrected data through /api/upload/input-template/edited.",
    )
    add_table(
        doc,
        ["Frontend area", "Technical role", "Backend APIs"],
        [
            ["UploadPage + ImportPreviewGrid", "File upload, row issue review, in-app correction, readiness display.", "Upload, Validation, Sessions"],
            ["DatabasePage", "CRUD and Excel upload/download for master data and lab requirements.", "Database"],
            ["SoftConstraintsPage", "Rank active soft preferences and trigger generation.", "Soft constraints, Schedules"],
            ["TimetableReviewPage + TimetableGrid + ConflictTable", "Visual timetable review, filters, move controls, conflict table, quick fixes.", "Schedules, Rooms, Timeslots, Violations"],
            ["ExportPage", "Load latest schedule and expose CSV/XLSX download links only when allowed.", "Schedules latest, Export"],
            ["DashboardPage", "Summarize readiness, latest run, conflicts, availability, and export shortcut.", "Dashboard, Availability, Constraint insights, Export"],
        ],
        [2250, 4300, 2810],
        font_size=8.7,
    )


def add_testing_future(doc):
    doc.add_heading("Testing, Verification, Limitations, and Future Work", level=1)
    add_para(
        doc,
        "The project should be verified at three levels: unit/service tests for backend logic, frontend build/lint checks for UI integration, and manual demo scripts that exercise the full workflow from upload to export.",
    )
    add_table(
        doc,
        ["Verification area", "What to check", "Why"],
        [
            ["Import tests", "Column aliases, two-sheet templates, preview rows, edited-row import, duplicate requirement IDs.", "Bad input handling is the first failure point in real use."],
            ["Validation tests", "Reference lookup, co-teachers, fixed time slots, feasible room checks, delivery/campus compatibility.", "Prevents unschedulable data from reaching the solver."],
            ["Solver tests", "Exact-once assignment, hard no-overlap, soft weights, no-candidate handling, relaxed/fallback behaviour.", "Confirms the mathematical model matches timetable rules."],
            ["Constraint audit tests", "Room/staff/group overlaps, capacity, delivery mismatch, soft warnings.", "Ensures review/export state is trustworthy."],
            ["Frontend checks", "Upload/edit flow, soft ranking, generate action, timetable review, manual move, export lock/unlock.", "Confirms users can operate the system end to end."],
        ],
        [1900, 4500, 2960],
        font_size=8.7,
    )
    add_heading = doc.add_heading
    add_heading("Known Limitations", level=2)
    add_bullets(
        doc,
        [
            "No authentication or role-based access control is currently described in the implementation.",
            "Imports replace the current requirement set and related schedule state rather than merging multiple active scenarios.",
            "The local SQLite setup is appropriate for prototype/demo use; production deployment would require migration planning and stronger backup/concurrency controls.",
            "The solver can produce reviewable fallback schedules for over-constrained input, so the post-generation conflict audit and export gate remain essential.",
            "Advanced institutional timetable rules, such as multi-campus travel times or official SIS sync, would require additional data fields and constraints.",
        ],
    )
    add_heading("Future Technical Work", level=2)
    add_bullets(
        doc,
        [
            "Add user accounts, roles, audit trails, and approval workflow for timetable changes.",
            "Support named scenarios so teams can compare multiple imported requirement sets without overwriting the current one.",
            "Add solver explanation summaries that show which soft preferences drove each placement.",
            "Introduce production database migrations and environment-specific configuration.",
            "Expand export/import adapters for official institutional systems if their final schemas differ from the current system-template file.",
        ],
    )


def add_appendix(doc):
    doc.add_heading("Appendix A: Important File Map", level=1)
    add_table(
        doc,
        ["Path", "Purpose"],
        [
            ["backend/app/main.py", "FastAPI app startup, CORS, router registration, health check."],
            ["backend/app/routes/upload_routes.py", "Workbook upload, preview, edited import, sample loaders."],
            ["backend/app/routes/data_routes.py", "Dashboard, reference reads, sessions CRUD, availability, constraint insights."],
            ["backend/app/routes/database_routes.py", "Reference database CRUD and Excel upload/download."],
            ["backend/app/routes/soft_constraint_routes.py", "Soft-constraint ranking API."],
            ["backend/app/routes/schedule_routes.py", "Generate, latest/run views, compare, manual move, quick fixes, recheck, explanations, violations."],
            ["backend/app/routes/export_routes.py", "CSV/XLSX download with hard-conflict gate."],
            ["backend/app/services/import_service.py", "Workbook parsing, alias normalization, preview and final import."],
            ["backend/app/services/requirement_input_service.py", "Input validation, reference resolution, session creation, co-teacher handling."],
            ["backend/app/services/scheduling_rules.py", "Shared candidate slot/room compatibility rules."],
            ["backend/app/services/schedule_service.py", "Generation orchestration and schedule persistence."],
            ["backend/app/solver/model_builder.py", "CP-SAT variables, hard constraints, soft penalties, objective."],
            ["backend/app/solver/cp_sat_solver.py", "Solver execution, strict/relaxed/greedy fallback strategy."],
            ["backend/app/solver/result_parser.py", "Converts selected solver variables into rows for persistence."],
            ["backend/app/services/constraint_service.py", "Post-generation hard/soft audit and violation storage."],
            ["backend/app/services/export_service.py", "Review rows and system-template CSV/XLSX export."],
            ["frontend/src/api/client.ts", "Central typed API wrapper and exported endpoint functions."],
            ["frontend/src/pages and frontend/src/components", "Workflow pages and reusable upload/review/export UI components."],
        ],
        [3600, 5760],
        font_size=8.3,
    )
    doc.add_heading("Appendix B: API Summary", level=1)
    add_table(
        doc,
        ["Workflow", "Frontend function", "Backend endpoint"],
        [
            ["Upload workbook", "uploadTemplate(files)", "POST /api/upload/input-template"],
            ["Apply edited import rows", "importEditedTemplateRows(rows)", "POST /api/upload/input-template/edited"],
            ["Read validation", "getValidation()", "GET /api/validation/latest"],
            ["Generate schedule", "generateSchedule()", "POST /api/schedules/generate"],
            ["Read latest schedule", "getLatestSchedule()", "GET /api/schedules/latest"],
            ["Move session", "moveScheduledSession(run, session, data)", "PUT /api/schedules/{run}/sessions/{session}"],
            ["Suggest fixes", "suggestScheduleFixes(run, data)", "POST /api/schedules/{run}/suggest-fixes"],
            ["Recheck run", "recheckSchedule(run)", "POST /api/schedules/{run}/recheck"],
            ["Download export", "exportUrl(run, 'csv'|'xlsx')", "GET /api/export/{run}/csv or /xlsx"],
            ["Manage soft priorities", "get/updateSoftConstraintPriorities()", "GET/PUT /api/soft-constraints"],
            ["Manage master data", "get/create/update/deleteDatabaseRow()", "GET/POST/PUT/DELETE /api/database/{type}"],
        ],
        [2150, 3550, 3660],
        font_size=8.3,
    )


def build():
    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_information(doc)
    add_contents(doc)
    add_architecture(doc)
    add_stack(doc)
    add_api_integration(doc)
    add_database_model(doc)
    add_import_pipeline(doc)
    add_requirement_validation(doc)
    add_schedule_service(doc)
    add_solver_model(doc)
    add_constraints(doc)
    add_solver_execution(doc)
    add_constraint_audit(doc)
    add_review_and_export(doc)
    add_frontend_workflow(doc)
    add_testing_future(doc)
    add_appendix(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
